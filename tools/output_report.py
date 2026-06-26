"""output_report.py - 文献综述报告输出工具（Phase 8）

支持三种输出格式：
1. Markdown (.md)   — 完整文献综述报告
2. Excel (.xlsx)    — 文献摘要表格
3. Word (.docx)     — 完整文献综述报告（Word格式）

用法:
    python tools/output_report.py \
        --input sorted.json \
        --format markdown \
        --topic "大语言模型在医疗诊断中的应用" \
        --date "2026-06-26"

    python tools/output_report.py \
        --input sorted.json \
        --format excel \
        --topic "LLM in Medical Diagnosis"

    python tools/output_report.py \
        --input sorted.json \
        --format all \
        --topic "大语言模型" \
        --overview "本文综述了..."

输入 (papers.json): 论文数组，每条含 title/authors/year/doi/abstract/source/compliant/tier/cited_by_count
输出: 按 format 参数生成对应文件
"""

import argparse
import json
import os
import sys
from datetime import date as dt_date


# ───────────────────────── 工具函数 ─────────────────────────


def _fmt_authors(authors: str | list | None) -> str:
    """格式化作者列表：取前 3 位，超过用"等"省略"""
    if not authors:
        return "-"
    if isinstance(authors, list):
        names = [str(a) for a in authors if a]
    elif isinstance(authors, str):
        names = [a.strip() for a in authors.split(",") if a.strip()]
    else:
        return "-"
    if len(names) <= 3:
        return ", ".join(names)
    return ", ".join(names[:3]) + ", 等"


def _fmt_doi(doi: str | None) -> str:
    """格式化 DOI"""
    if not doi:
        return "-"
    return doi.replace("https://doi.org/", "")


def _fmt_venue(paper: dict) -> str:
    """提取期刊/会议名称"""
    venue = paper.get("venue")
    if venue:
        return venue
    source = paper.get("source", "")
    if "arXiv" in source:
        return "arXiv preprint"
    return "-"


def _fmt_citations(count) -> str:
    """格式化引用量"""
    if count is None:
        return "-"
    try:
        return str(int(count))
    except (TypeError, ValueError):
        return "-"


def _fmt_compliant(val) -> str:
    """格式化合规标记"""
    if val:
        return "是"
    return "否"


def _fmt_tier(tier: str | None) -> str:
    """格式化等级"""
    return tier or "-"


def _fmt_url(paper: dict) -> str:
    """构建文章链接"""
    doi = paper.get("doi")
    if doi:
        return f"https://doi.org/{_fmt_doi(doi)}"
    pdf_url = paper.get("pdf_url")
    if pdf_url:
        return pdf_url
    return "-"


def _load_papers(path: str) -> list[dict]:
    """加载论文 JSON 文件，兼容 {papers:[...]} 和 [...] 两种格式"""
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if isinstance(data, dict) and "papers" in data:
        return data["papers"]
    if isinstance(data, list):
        return data
    raise ValueError(f"无法解析 {path}：期望论文数组或含 papers 字段的对象")


def _suggest_filename(topic: str, date_str: str, fmt: str = "md") -> str:
    """根据格式生成对应的文件名"""
    safe_topic = "".join(c if c.isalnum() or c in " _-" else "_" for c in topic)
    safe_topic = safe_topic.strip().replace(" ", "_")[:60]
    if fmt == "excel":
        return f"{safe_topic}文献信息表_{date_str}.xlsx"
    ext = ".docx" if fmt == "word" else ".md"
    return f"{safe_topic}文献综述_{date_str}{ext}"


# ─────────────────── Markdown 输出 ───────────────────


def generate_markdown(papers: list[dict], topic: str, date_str: str,
                      overview: str = "") -> str:
    """生成 Markdown 格式的文献综述报告"""
    total = len(papers)
    premium = sum(1 for p in papers if p.get("compliant"))
    baseline = total - premium

    lines = []
    # 标题
    lines.append(f"# {topic} 文献综述\n")
    lines.append(f"> 检索时间：{date_str}")
    lines.append(f"> 数据来源：arXiv + OpenAlex")
    lines.append(f"> 文献总数：{total} 篇（符合等级要求 {premium} 篇 / 其他 {baseline} 篇）")
    lines.append("")
    lines.append("---\n")

    # 领域概述
    lines.append("## 一、领域概述\n")
    if overview:
        lines.append(f"{overview}\n")
    else:
        lines.append("（由 AI 在 Phase 7 根据摘要内容自动生成）\n")
    lines.append("---\n")

    # 文献表格


    # GB/T 7714 参考文献
    lines.append("## 二、参考文献\n")
    for i, p in enumerate(papers, 1):
        authors = _fmt_authors(p.get("authors"))
        title = p.get("title", "")
        venue = _fmt_venue(p)
        year = p.get("year", "")
        doi = _fmt_doi(p.get("doi"))
        paper_type = p.get("type", "")
        if not paper_type:
            if p.get("source") in ("arXiv",):
                paper_type = "J"
            elif "Conference" in str(p.get("type", "")):
                paper_type = "C"
            else:
                paper_type = "J"
        type_tag = "[J]" if "J" in str(paper_type) or "Journal" in str(p.get("type", "")) else "[C]"
        ref = f"{i}. {authors}. {title}{type_tag}. {venue}, {year}."
        if doi and doi != "-":
            ref += f" DOI: {doi}"
        lines.append(ref)

    return "\n".join(lines)


# ─────────────────── Excel 输出 ───────────────────


def generate_excel(papers: list[dict], topic: str, date_str: str, output_dir: str) -> str:
    """生成 Excel 摘要表格"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "文献摘要"

    # 列头
    headers = ["序号", "标题", "作者", "年份", "来源", "DOI", "文章链接", "期刊/会议",
               "引用量", "等级", "符合等级要求", "摘要"]
    header_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border

    # 数据行
    wrap_align = Alignment(vertical="top", wrap_text=True)
    for i, p in enumerate(papers, 1):
        row_data = [
            i,
            p.get("title", ""),
            _fmt_authors(p.get("authors")),
            p.get("year", ""),
            p.get("source", ""),
            _fmt_doi(p.get("doi")),
            _fmt_url(p),
            _fmt_venue(p),
            _fmt_citations(p.get("cited_by_count")),
            _fmt_tier(p.get("tier")),
            _fmt_compliant(p.get("compliant")),
            p.get("abstract", ""),
        ]
        for col, val in enumerate(row_data, 1):
            cell = ws.cell(row=i + 1, column=col, value=val)
            cell.alignment = wrap_align
            cell.border = thin_border

    # 列宽
    col_widths = [6, 40, 20, 8, 12, 30, 35, 25, 10, 18, 12, 60]
    for col, w in enumerate(col_widths, 1):
        ws.column_dimensions[chr(64 + col)].width = w

    filename = _suggest_filename(topic, date_str, "excel")
    filepath = os.path.join(output_dir, filename)
    wb.save(filepath)
    return filepath


# ─────────────────── Word 输出 ───────────────────


def generate_word(papers: list[dict], topic: str, date_str: str,
                  overview: str, output_dir: str) -> str:
    """生成 Word 格式的文献综述报告"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 标题
    title = doc.add_heading(f"{topic} 文献综述", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元信息
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"检索时间：{date_str}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    meta.add_run("\n")
    run = meta.add_run("数据来源：arXiv + OpenAlex")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    total = len(papers)
    premium = sum(1 for p in papers if p.get("compliant"))
    baseline = total - premium
    meta.add_run("\n")
    run = meta.add_run(f"文献总数：{total} 篇（符合等级要求 {premium} 篇 / 其他 {baseline} 篇）")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # 领域概述
    doc.add_heading("一、领域概述", level=1)
    if overview:
        doc.add_paragraph(overview)
    else:
        doc.add_paragraph("（由 AI 在 Phase 7 根据摘要内容自动生成）")

    # 文献表格

    doc.add_heading("二、参考文献", level=1)
    for i, p in enumerate(papers, 1):
        authors = _fmt_authors(p.get("authors"))
        title = p.get("title", "")
        venue = _fmt_venue(p)
        year = p.get("year", "")
        doi = _fmt_doi(p.get("doi"))
        type_tag = "[J]"
        if "Conference" in str(p.get("type", "")):
            type_tag = "[C]"
        elif p.get("source") == "arXiv":
            type_tag = "[J]"
        ref_text = f"{i}. {authors}. {title}{type_tag}. {venue}, {year}."
        if doi and doi != "-":
            ref_text += f" DOI: {doi}"
        doc.add_paragraph(ref_text, style="List Number")

    filename = _suggest_filename(topic, date_str, "word")
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


# ─────────────────── 主入口 ───────────────────


def main():
    parser = argparse.ArgumentParser(
        description="output_report.py - 文献综述报告输出工具（Phase 8）"
    )
    parser.add_argument("--input", required=True,
                        help="输入论文 JSON 文件路径")
    parser.add_argument("--format", required=True,
                        choices=["markdown", "excel", "word", "all"],
                        help="输出格式：markdown / excel / word / all")
    parser.add_argument("--output-dir", default=".",
                        help="输出目录（默认当前目录）")
    parser.add_argument("--topic", default="文献综述",
                        help="研究主题，用于文件名和报告标题")
    parser.add_argument("--date", default=dt_date.today().isoformat(),
                        help="日期字符串（默认今天）")
    parser.add_argument("--overview", default="",
                        help="领域概述文本（可选，用于报告）")
    args = parser.parse_args()

    # 加载数据
    try:
        papers = _load_papers(args.input)
    except Exception as e:
        print(f"[ERROR] 加载输入文件失败: {e}", file=sys.stderr)
        sys.exit(1)

    if not papers:
        print("[ERROR] 论文列表为空", file=sys.stderr)
        sys.exit(1)

    os.makedirs(args.output_dir, exist_ok=True)
    generated = []

    # Markdown
    if args.format in ("markdown", "all"):
        md_content = generate_markdown(papers, args.topic, args.date, args.overview)
        md_file = _suggest_filename(args.topic, args.date, "markdown")
        md_path = os.path.join(args.output_dir, md_file)
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        print(f"[INFO] Markdown 报告已写入: {os.path.abspath(md_path)}", file=sys.stderr)
        generated.append(md_path)

    # Excel
    if args.format in ("excel", "all"):
        xlsx_path = generate_excel(papers, args.topic, args.date, args.output_dir)
        print(f"[INFO] Excel 摘要表已写入: {os.path.abspath(xlsx_path)}", file=sys.stderr)
        generated.append(xlsx_path)

    # Word
    if args.format in ("word", "all"):
        docx_path = generate_word(papers, args.topic, args.date, args.overview, args.output_dir)
        print(f"[INFO] Word 报告已写入: {os.path.abspath(docx_path)}", file=sys.stderr)
        generated.append(docx_path)

    # 输出汇总
    print(json.dumps({"generated": generated}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
